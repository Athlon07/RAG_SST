"""Load text content from various document types into a common structure.

Every loader returns a list of "sections":
    [{"location": str, "text": str}, ...]

`location` is a human-readable label for where the text came from - a page
number for PDFs, a sheet name for spreadsheets, a heading for Word/HTML
sections, etc. Everything downstream (chunking, vector store, citations)
just treats it as an opaque label, so adding a new file type only means
adding a loader function here and one line in `load_document`.
"""
import os
import xml.etree.ElementTree as ET
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import openpyxl
from bs4 import BeautifulSoup


def load_pdf(file_path: str):
    """Extract text from each page of a PDF using PyMuPDF.

    PyMuPDF (fitz) is used instead of pypdf because it's substantially faster
    at text extraction, which matters most for larger documents / batches of
    files. Pages with no extractable text (e.g. scanned/image-only pages)
    are skipped - those would need OCR.
    """
    sections = []
    with fitz.open(file_path) as doc:
        for i, page in enumerate(doc):
            text = page.get_text().strip()
            if text:
                sections.append({"location": f"Page {i + 1}", "text": text})
    return sections


def load_docx(file_path: str):
    """Extract text from a Word document, grouped by heading. Each heading
    (Heading 1/2/3 style) starts a new section; body paragraphs under it are
    grouped together. Tables are extracted separately as their own sections."""
    doc = DocxDocument(file_path)
    sections = []
    current_heading = "Introduction"
    current_texts = []

    def flush():
        if current_texts:
            sections.append({"location": current_heading, "text": "\n".join(current_texts)})

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            flush()
            current_heading = text
            current_texts = []
        else:
            current_texts.append(text)
    flush()

    for t_idx, table in enumerate(doc.tables, start=1):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        if rows_text:
            sections.append({"location": f"Table {t_idx}", "text": "\n".join(rows_text)})

    return sections


def load_xlsx(file_path: str):
    """Extract each worksheet as its own section. The first non-empty row is
    treated as a header row; every subsequent row is turned into
    'header: value' pairs so the embedding model gets column context, not
    just a raw table dump."""
    wb = openpyxl.load_workbook(file_path, data_only=True)
    sections = []

    for sheet in wb.worksheets:
        rows_text = []
        headers = None
        for row in sheet.iter_rows(values_only=True):
            values = ["" if v is None else str(v) for v in row]
            if not any(v.strip() for v in values):
                continue
            if headers is None:
                headers = values
                continue
            pairs = [f"{h}: {v}" for h, v in zip(headers, values) if v.strip()]
            if pairs:
                rows_text.append(", ".join(pairs))
        if rows_text:
            sections.append({"location": f"Sheet: {sheet.title}", "text": "\n".join(rows_text)})

    return sections


def load_html(file_path: str):
    """Extract text from an HTML file, split into sections by heading tags
    (h1/h2/h3) so retrieval can point back to a rough section name."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    for tag in soup(["script", "style"]):
        tag.decompose()

    body = soup.body or soup
    sections = []
    current_heading = "Document"
    current_texts = []

    def flush():
        if current_texts:
            text = "\n".join(t for t in current_texts if t)
            if text.strip():
                sections.append({"location": current_heading, "text": text})

    for el in body.find_all(["h1", "h2", "h3", "p", "li", "td"]):
        text = el.get_text(" ", strip=True)
        if not text:
            continue
        if el.name in ("h1", "h2", "h3"):
            flush()
            current_heading = text
            current_texts = []
        else:
            current_texts.append(text)
    flush()

    if not sections:
        text = soup.get_text("\n", strip=True)
        if text:
            sections.append({"location": "Document", "text": text})

    return sections


def load_xml(file_path: str):
    """Generic XML text extraction, one section per element that directly
    contains text, labeled by its tag path. Works well for document-like XML
    (articles, feeds, docbook); for deeply nested data XML you may want a
    schema-specific loader instead."""
    tree = ET.parse(file_path)
    root = tree.getroot()
    sections = []

    def walk(el, path):
        text = (el.text or "").strip()
        if text:
            sections.append({"location": path, "text": text})
        for child in el:
            walk(child, f"{path}/{child.tag}")

    walk(root, root.tag)
    return sections


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
    ".xlsm": load_xlsx,
    ".html": load_html,
    ".htm": load_html,
    ".xml": load_xml,
}


def load_document(file_path: str):
    """Dispatch to the right loader based on file extension.

    Returns list[dict]: [{"location": str, "text": str}, ...]
    Raises ValueError for unsupported extensions.
    """
    ext = os.path.splitext(file_path)[1].lower()
    loader = LOADERS.get(ext)
    if loader is None:
        supported = ", ".join(sorted(LOADERS.keys()))
        raise ValueError(f"Unsupported file type '{ext}'. Supported types: {supported}")
    return loader(file_path)


def load_documents_parallel(file_paths, max_workers: int = 4):
    """Load multiple files concurrently using a thread pool.

    Text extraction (PyMuPDF, python-docx, openpyxl, BeautifulSoup) is
    largely I/O-bound plus C-extension work that releases the GIL, so a
    thread pool gives a real speedup here without the complexity of
    multiprocessing (no pickling of results/loader functions needed).

    Returns:
        dict[str, list]: {file_path: sections} for files that loaded successfully.
        errors: dict[str, str] mapping file_path -> error message, for files
        that failed to load (unsupported type, corrupt file, etc.) - callers
        can decide whether to surface these or skip silently.
    """
    results = {}
    errors = {}

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_path = {executor.submit(load_document, fp): fp for fp in file_paths}
        for future in as_completed(future_to_path):
            fp = future_to_path[future]
            try:
                results[fp] = future.result()
            except Exception as e:
                errors[fp] = str(e)

    return results, errors
